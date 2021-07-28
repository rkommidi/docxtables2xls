#!/usr/bin/env python

from docx.api import Document
import pandas as pd

path = "/Users/Raghav/Work/Projects/ReportConversion/in/Test2.docx"
document = Document(path)
output_path = "/Users/Raghav/Work/Projects/ReportConversion/out"
writer = pd.ExcelWriter('{}/1scripts.xlsx'.format(output_path), engine='xlsxwriter')

#Results Summary
table = document.tables[1]
data = []
keys = []
values = []

for row in table.rows:
    for index, cell in enumerate(row.cells):
        if index == 0 or index == 2:
            keys.append(cell.text)
        else:
            values.append(cell.text)

headers = ['Result Summary', 'Run']
for index in range(len(keys)):
    row_data = dict(zip(headers, (keys[index], values[index])))
    data.append(row_data)
    print(row_data)
rdf = pd.DataFrame(data)
rdf.to_excel(writer, sheet_name='Summary',index=False)

#Statistics Summary
table = document.tables[2]
data = []
keys = []
values = []

for row in table.rows:
    for index, cell in enumerate(row.cells):
        if index == 0 or index == 3:
            keys.append(cell.text)
        elif index == 1 or index == 4:
            values.append(cell.text)

headers = ['Statistics Summary', 'Run']
for index in range(len(keys)):
    row_data = dict(zip(headers, (keys[index], values[index])))
    data.append(row_data)
    print(row_data)
sdf = pd.DataFrame(data)
sdf.to_excel(writer, sheet_name='Summary',index=False, columns=headers, startrow=20, startcol=0)

no_of_rows = 0
cols = ['Transaction Name', 'Avg', 'Avg-90%', 'Count', 'Err', 'Std Dev', 'SLA Profile']
for table_num in range(16,17):
    print(table_num)
    table = document.tables[table_num]
    data = []
    keys = None
    row_data = None
    flag = 0
    for j, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if j == 0:
            keys = list(text)
            keys.insert(0,'Transaction Name')
            continue
        if flag == 0:
            text_list = tuple(text)
            if text_list[0][:4] == 'MRSi':
                flag = 1
                transaction_name = text_list[0]
                continue
        if flag == 1:
            flag = 0
            text_list = list(text)
            text_list.insert(0,transaction_name)
            row_data = dict(zip(keys, text_list))
            data.append(row_data)
    #        print(row_data)
    df = pd.DataFrame(data)
    df.to_excel(writer, sheet_name='Response Times', columns=cols, index=False, startrow=no_of_rows, startcol=0)
    no_of_rows += df.shape[0] + 3

#Set Column Width
writer.sheets['Summary'].set_column('A:A',50)
writer.sheets['Response Times'].set_column('A:A',50)

writer.save()
