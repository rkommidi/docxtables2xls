#!C:\Users\ragha\AppData\Local\Programs\Python\Python39\python.exe

from datetime import datetime
import sys
import re
import datetime
from docx.api import Document
from docx.shared import Cm
import pandas as pd

#################################################################
# Command Line Argumnets
# python.exe .\neo_to_lr.py Policy_4_2_report 10 10
#################################################################
input_file = sys.argv[1]
no_of_users = sys.argv[2]
no_of_tables = sys.argv[3]
print(input_file,no_of_users, no_of_tables)

base_path = "C:\\Users\\ragha\\OneDrive\\Desktop\\padhu\\Test\\"

path = "{}{}.docx".format(base_path, input_file)
document = Document(path)
output_path = base_path
writer = pd.ExcelWriter('{}/{}_output.xlsx'.format(output_path,input_file), engine='xlsxwriter')

word_document = Document()
document_name = '{}_output'.format(input_file)


#Results Summary
table = document.tables[1]
data = []
keys = []
values = []

summary_mapping = {
                     'Name': 'Run ID',                              #covert to DDHHSS format                    
                     'Status': 'Status',
                     'Duration': 'Duration',
                     'Total request errors': 'Script Errors',
                     'Start date': 'Date',                          # strip Date
                     'End date': 'Time',                            # fetch time form start date
                     'Filters': 'Time Range',                        # strip time range
                     'LG Hosts': 'Percentile Values',               # hardcord to 85 percentile
                     'Average throughput': 'Average Throughput',
                     'Total throughput': 'Total Throughput',
                     'Average requests/s': 'Average Hits',
                     'Total requests': 'Total Hits',
                     'Total pages': 'Passed Transactions',
                     'Total action errors': 'Failed Transactions',  # replace with Total request errors
                     'Total users launched': 'Total WEB Vusers'     # pass from command line
                  }              

trans_mapping = {
                    'Transaction Name': 'Transaction',
                    'Min': 'MIN',
                    'Avg': 'AVG',
                    'Max': 'MAX',
                    'Count': 'PASSED',
                    'Err': 'FAILED',
                    '% of Err': '% of Err',
                    'Perc 80': '85th',
                    'Perc 90': '90th',
                    'Perc 95': '95th',
                    'Std Dev': 'STD',
                    'Avg-90%': 'Avg-90%',
                    'SLA Profile': 'SLA Profile'
                }
for row in table.rows:
    for index, cell in enumerate(row.cells):
        if index == 0 or index == 2:
            keys.append(cell.text)
        else:
            values.append(cell.text)

headers = ['Result Summary', 'Run']
summary_hash = {}

for index in range(len(keys)):
    row_data = dict(zip(headers, (keys[index], values[index])))
    data.append(row_data)
    summary_hash[keys[index]] = values[index]
    #print(row_data)

rdf = pd.DataFrame(data)
rdf.to_excel(writer, sheet_name='Summary',index=False)

#Statistics Summary
table = document.tables[2]
data = []
keys = []
values = []

for row in table.rows:
    for index, cell in enumerate(row.cells):
        if index == 0 or index == 3:
            keys.append(cell.text)
        elif index == 1 or index == 4:
            values.append(cell.text)

headers = ['Statistics Summary', 'Run']
for index in range(len(keys)):
    row_data = dict(zip(headers, (keys[index], values[index])))
    data.append(row_data)
    summary_hash[keys[index]] = values[index]
    print(row_data)

# Convert to LoadRunner Word Format
word_table = word_document.add_table(0, 0)
word_table.style = 'TableGrid'
first_column_width = 15
second_column_with = 15
word_table.add_column(Cm(first_column_width))
word_table.add_column(Cm(second_column_with))
word_table_index = 0
for mapping_key in summary_mapping.keys():
    word_table.add_row()
    row = word_table.rows[word_table_index]
    word_table_index += 1
    summary_key = str(summary_mapping[mapping_key])
    if summary_key == 'Percentile Values':
        summary_val = '85'
    elif summary_key == 'Date':
        summary_val = str(summary_hash[mapping_key]).split(",")[0:2]
    elif summary_key == 'Time':
        summary_val = str(summary_hash["Start date"]).split(",")[2]
    elif summary_key == 'Time Range':
        filters = str(summary_hash[mapping_key])
        #print(filters.split(" "))
        #summary_val = filters.split(" ")[5].split(".")[0] + '-' + filters.split(" ")[8].split(".")[0]
        res = re.findall(r"(\d{2}:\d{2}:\d{2}).*(\d{2}:\d{2}:\d{2})", filters) 
        summary_val = "-".join(res[0])
    elif summary_key == 'Total action errors':
        summary_val = str(summary_hash["Total request errors"])      
    elif summary_key == 'Total WEB Vusers':
        summary_val = str(no_of_users) 
    elif summary_key == 'Run ID':
        #date_res = re.findall(r"(\d{2}):(\d{2}).*(\d{2}).*(\w{3]).*(\d{4})}) ",summary_hash[mapping_key])
        #summary_val = str(date_res[0][0])+str(date_res[0][1])
        summary_val = str(summary_hash[mapping_key])
    else:
        summary_val = str(summary_hash[mapping_key])
    
    row.cells[0].text = summary_key
    row.cells[1].text = summary_val

sdf = pd.DataFrame(data)
sdf.to_excel(writer, sheet_name='Summary',index=False, columns=headers, startrow=20, startcol=0)

no_of_rows = 0
cols = ['Transaction Name', 'Min', 'Avg', 'Max', 'Count', 'Err', '% of Err', 'Perc 80', 'Perc 90', 'Perc 95', 'Std Dev', 'Avg-90%', 'SLA Profile']
script_index = 31
word_document.add_heading("Transaction", 1)
for table_num in range(16, 16 + int(no_of_tables)):
    print(table_num)
    print(document.paragraphs[script_index].text)
    word_document.add_heading(str(document.paragraphs[script_index].text),2)
    script_index += 1
    trans_table = word_document.add_table(0, 0)
    trans_table.style = 'TableGrid'
    trans_table_index = 0
    for col in cols:
        trans_table.add_column(Cm(1))
    
    trans_table.add_row()
    row = trans_table.rows[trans_table_index]
    trans_table_index += 1
    for col_index, col in enumerate(cols):
        row.cells[col_index].text = str(trans_mapping[col])

    table = document.tables[table_num]
    data = []
    keys = None
    row_data = None
    flag = 0
    for j, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if j == 0:
            keys = list(text)
            keys.insert(0,'Transaction Name')
            continue
        if flag == 0:
            text_list = tuple(text)
            if text_list[0][:1] in ('T', 'M', 'S'):
                flag = 1
                transaction_name = text_list[0]
                continue
        if flag == 1:
            flag = 0
            text_list = list(text)
            text_list.insert(0,transaction_name)
            row_data = dict(zip(keys, text_list))
            data.append(row_data)
            #print(row_data)
            trans_table.add_row()
            row = trans_table.rows[trans_table_index]
            trans_table_index += 1
            for col_index, col in enumerate(cols):
                row.cells[col_index].text = str(row_data[col])
    df = pd.DataFrame(data)
    df.to_excel(writer, sheet_name='Response Times', columns=cols, index=False, startrow=no_of_rows, startcol=0)
    no_of_rows += df.shape[0] + 3

#Set Column Width
writer.sheets['Summary'].set_column('A:A',50)
writer.sheets['Response Times'].set_column('A:A',50)

#for i,para in enumerate(document.paragraphs):
#    print(i,para.text)
writer.save()

word_document.add_page_break()

word_document.save('{}/{}.docx'.format(output_path,document_name))